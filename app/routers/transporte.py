from __future__ import annotations
import io
import os
from datetime import datetime, date
from fastapi import APIRouter, Request, Form, Depends, HTTPException
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import or_
from app.database import get_db
from app.deps import get_current_user
from app import models_transporte as mt
from app.templates import templates

# FORCE REDEPLOY - Railway cache fix
router = APIRouter(prefix="/transporte")


def _require_access(current_user=Depends(get_current_user)):
    if current_user.role == "operador":
        raise HTTPException(status_code=403, detail="Sin acceso al módulo de Transporte.")
    return current_user


def _parse_date(s: str):
    if not s or not s.strip():
        return None
    try:
        return datetime.strptime(s.strip(), "%Y-%m-%d").date()
    except ValueError:
        return None


# ── NÓMINA MADRE ───────────────────────────────────────────────────────────────

@router.get("/nomina", response_class=HTMLResponse)
async def nomina_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    q = request.query_params.get("q", "").strip()

    query = db.query(mt.TransporteNomina).filter(mt.TransporteNomina.deleted_at == None)

    if q:
        query = query.filter(
            or_(
                mt.TransporteNomina.empresa.ilike(f"%{q}%"),
                mt.TransporteNomina.nombre_chofer.ilike(f"%{q}%"),
                mt.TransporteNomina.dni.ilike(f"%{q}%"),
                mt.TransporteNomina.patente_camion.ilike(f"%{q}%"),
                mt.TransporteNomina.patente_acoplado.ilike(f"%{q}%"),
            )
        )

    nomina = (
        query
        .order_by(
            mt.TransporteNomina.empresa.asc(),
            mt.TransporteNomina.nombre_chofer.asc(),
        )
        .all()
    )

    return templates.TemplateResponse(request, "transporte/nomina_list.html", {
        "request": request,
        "user": current_user,
        "nomina": nomina,
        "q": q,
    })


@router.get("/nomina/new", response_class=HTMLResponse)
async def nomina_new_form(
    request: Request,
    current_user=Depends(_require_access),
):
    return templates.TemplateResponse(request, "transporte/nomina_form.html", {
        "request": request,
        "user": current_user,
        "item": None,
        "errors": [],
    })


@router.post("/nomina/new", response_class=HTMLResponse)
async def nomina_create(
    request: Request,
    empresa: str          = Form(""),
    nombre_chofer: str    = Form(""),
    dni: str              = Form(""),
    marca_camion: str     = Form(""),
    patente_camion: str   = Form(""),
    patente_acoplado: str = Form(""),
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    errors = []
    if not empresa.strip():
        errors.append("La empresa es requerida.")
    if not nombre_chofer.strip():
        errors.append("El nombre del chofer es requerido.")

    if errors:
        return templates.TemplateResponse(request, "transporte/nomina_form.html", {
            "request": request,
            "user": current_user,
            "item": None,
            "errors": errors,
            "form": {
                "empresa": empresa,
                "nombre_chofer": nombre_chofer,
                "dni": dni,
                "marca_camion": marca_camion,
                "patente_camion": patente_camion,
                "patente_acoplado": patente_acoplado,
            },
        }, status_code=422)

    item = mt.TransporteNomina(
        empresa          = empresa.strip(),
        nombre_chofer    = nombre_chofer.strip(),
        dni              = dni.strip() or None,
        marca_camion     = marca_camion.strip() or None,
        patente_camion   = patente_camion.strip().upper() or None,
        patente_acoplado = patente_acoplado.strip().upper() or None,
    )
    db.add(item)
    db.commit()
    return RedirectResponse("/transporte/nomina", status_code=303)


@router.get("/nomina/{item_id}/edit", response_class=HTMLResponse)
async def nomina_edit_form(
    item_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    item = db.query(mt.TransporteNomina).filter(
        mt.TransporteNomina.id == item_id,
        mt.TransporteNomina.deleted_at == None,
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="Registro no encontrado.")

    return templates.TemplateResponse(request, "transporte/nomina_form.html", {
        "request": request,
        "user": current_user,
        "item": item,
        "errors": [],
    })


@router.post("/nomina/{item_id}/edit", response_class=HTMLResponse)
async def nomina_update(
    item_id: int,
    request: Request,
    empresa: str          = Form(""),
    nombre_chofer: str    = Form(""),
    dni: str              = Form(""),
    marca_camion: str     = Form(""),
    patente_camion: str   = Form(""),
    patente_acoplado: str = Form(""),
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    item = db.query(mt.TransporteNomina).filter(
        mt.TransporteNomina.id == item_id,
        mt.TransporteNomina.deleted_at == None,
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="Registro no encontrado.")

    errors = []
    if not empresa.strip():
        errors.append("La empresa es requerida.")
    if not nombre_chofer.strip():
        errors.append("El nombre del chofer es requerido.")

    if errors:
        return templates.TemplateResponse(request, "transporte/nomina_form.html", {
            "request": request,
            "user": current_user,
            "item": item,
            "errors": errors,
            "form": {
                "empresa": empresa,
                "nombre_chofer": nombre_chofer,
                "dni": dni,
                "marca_camion": marca_camion,
                "patente_camion": patente_camion,
                "patente_acoplado": patente_acoplado,
            },
        }, status_code=422)

    item.empresa          = empresa.strip()
    item.nombre_chofer    = nombre_chofer.strip()
    item.dni              = dni.strip() or None
    item.marca_camion     = marca_camion.strip() or None
    item.patente_camion   = patente_camion.strip().upper() or None
    item.patente_acoplado = patente_acoplado.strip().upper() or None
    item.updated_at       = datetime.utcnow()
    db.commit()
    return RedirectResponse("/transporte/nomina", status_code=303)


@router.post("/nomina/{item_id}/delete")
async def nomina_delete(
    item_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    if current_user.role not in ("admin", "superadmin"):
        raise HTTPException(status_code=403, detail="Solo administradores pueden eliminar registros.")

    item = db.query(mt.TransporteNomina).filter(
        mt.TransporteNomina.id == item_id,
        mt.TransporteNomina.deleted_at == None,
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="Registro no encontrado.")

    item.deleted_at = datetime.utcnow()
    db.commit()
    return RedirectResponse("/transporte/nomina", status_code=303)


# ── HISTORIAL ─────────────────────────────────────────────────────────────────

@router.get("/historial", response_class=HTMLResponse)
async def historial_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    operativos = (
        db.query(mt.TransporteOperativo)
        .filter(mt.TransporteOperativo.deleted_at == None)
        .order_by(mt.TransporteOperativo.fecha_inicio.desc().nullslast(),
                  mt.TransporteOperativo.created_at.desc())
        .all()
    )

    return templates.TemplateResponse(request, "transporte/historial_list.html", {
        "request": request,
        "user": current_user,
        "operativos": operativos,
    })


@router.get("/historial/new", response_class=HTMLResponse)
async def historial_new_form(
    request: Request,
    current_user=Depends(_require_access),
):
    return templates.TemplateResponse(request, "transporte/historial_new.html", {
        "request": request,
        "user": current_user,
        "errors": [],
    })


@router.post("/historial/new", response_class=HTMLResponse)
async def historial_create(
    request: Request,
    nombre_barco: str = Form(""),
    producto: str     = Form(""),
    cliente: str      = Form(""),
    deposito: str     = Form(""),
    mercaderia_a_mover: str = Form(""),
    fecha_inicio: str = Form(""),
    fecha_fin: str    = Form(""),
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    errors = []
    if not nombre_barco.strip():
        errors.append("El nombre del barco es requerido.")

    if errors:
        return templates.TemplateResponse(request, "transporte/historial_new.html", {
            "request": request,
            "user": current_user,
            "errors": errors,
            "form": {
                "nombre_barco": nombre_barco,
                "producto": producto,
                "cliente": cliente,
                "deposito": deposito,
                "mercaderia_a_mover": mercaderia_a_mover,
                "fecha_inicio": fecha_inicio,
                "fecha_fin": fecha_fin,
            },
        }, status_code=422)

    op = mt.TransporteOperativo(
        nombre_barco  = nombre_barco.strip(),
        producto      = producto.strip() or None,
        cliente       = cliente.strip() or None,
        deposito      = deposito.strip() or None,
        fecha_inicio  = _parse_date(fecha_inicio),
        mercaderia_a_mover = mercaderia_a_mover.strip() or None,
        fecha_fin     = _parse_date(fecha_fin),
        created_by_id = current_user.id,
    )
    db.add(op)
    db.commit()
    db.refresh(op)
    return RedirectResponse(f"/transporte/historial/{op.id}/puerto", status_code=303)


@router.get("/historial/{op_id}", response_class=HTMLResponse)
async def historial_detail(
    op_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    op = (
        db.query(mt.TransporteOperativo)
        .options(
            joinedload(mt.TransporteOperativo.asignaciones)
            .joinedload(mt.TransporteOperativoAsignacion.assigned_by),
        )
        .filter(
            mt.TransporteOperativo.id == op_id,
            mt.TransporteOperativo.deleted_at == None,
        )
        .first()
    )
    if not op:
        raise HTTPException(status_code=404, detail="Operativo no encontrado.")

    # IDs ya asignados para no mostrarlos en el selector
    asignados_ids = {a.nomina_id for a in op.asignaciones if a.nomina_id}

    disponibles = (
        db.query(mt.TransporteNomina)
        .filter(
            mt.TransporteNomina.deleted_at == None,
            mt.TransporteNomina.id.notin_(asignados_ids) if asignados_ids else True,
        )
        .order_by(
            mt.TransporteNomina.empresa.asc(),
            mt.TransporteNomina.nombre_chofer.asc(),
        )
        .all()
    )

    return templates.TemplateResponse(request, "transporte/historial_detail.html", {
        "request": request,
        "user": current_user,
        "op": op,
        "disponibles": disponibles,
    })



@router.get("/historial/{op_id}/balanza", response_class=HTMLResponse)
async def historial_balanza(
    op_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    return await historial_detail(op_id, request, db, current_user)


@router.get("/historial/{op_id}/edit", response_class=HTMLResponse)
async def historial_edit_form(
    op_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    op = db.query(mt.TransporteOperativo).filter(
        mt.TransporteOperativo.id == op_id,
        mt.TransporteOperativo.deleted_at == None,
    ).first()
    if not op:
        raise HTTPException(status_code=404, detail="Operativo no encontrado.")

    return templates.TemplateResponse(request, "transporte/historial_edit.html", {
        "request": request,
        "user": current_user,
        "op": op,
        "errors": [],
    })


@router.post("/historial/{op_id}/edit", response_class=HTMLResponse)
async def historial_edit_save(
    op_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
    nombre_barco: str = Form(""),
    producto: str = Form(""),
    cliente: str = Form(""),
    deposito: str = Form(""),
    mercaderia_a_mover: str = Form(""),
    fecha_inicio: str = Form(""),
    fecha_fin: str = Form(""),
):
    op = db.query(mt.TransporteOperativo).filter(
        mt.TransporteOperativo.id == op_id,
        mt.TransporteOperativo.deleted_at == None,
    ).first()
    if not op:
        raise HTTPException(status_code=404, detail="Operativo no encontrado.")
    op.nombre_barco = nombre_barco
    op.producto = producto
    op.cliente = cliente
    op.deposito = deposito
    op.mercaderia_a_mover = mercaderia_a_mover
    from datetime import date as _date
    op.fecha_inicio = _date.fromisoformat(fecha_inicio) if fecha_inicio else None
    op.fecha_fin = _date.fromisoformat(fecha_fin) if fecha_fin else None
    db.commit()
    return RedirectResponse(url=f"/transporte/historial/{op_id}", status_code=303)


@router.post("/historial/{op_id}/delete", response_class=HTMLResponse)
async def historial_delete(
    op_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    op = db.query(mt.TransporteOperativo).filter(
        mt.TransporteOperativo.id == op_id,
        mt.TransporteOperativo.deleted_at == None,
    ).first()
    if not op:
        raise HTTPException(status_code=404, detail="Operativo no encontrado.")
    from datetime import datetime
    op.deleted_at = datetime.now()
    db.commit()
    return RedirectResponse(url="/transporte/historial", status_code=303)


@router.get("/historial/{op_id}/exportar-word-puerto")
async def historial_exportar_word_puerto(
    op_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    try:
        from docx import Document
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(
            status_code=501,
            detail="python-docx no está instalado.",
        )
    from pathlib import Path

    op = (
        db.query(mt.TransporteOperativo)
        .options(joinedload(mt.TransporteOperativo.asignaciones))
        .filter(
            mt.TransporteOperativo.id == op_id,
            mt.TransporteOperativo.deleted_at == None,
        )
        .first()
    )
    if not op:
        raise HTTPException(status_code=404, detail="Operativo no encontrado.")

    MESES = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
        5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
        9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
    }
    hoy = date.today()
    fecha_str = (
        f"San Nicolás de los Arroyos, "
        f"{hoy.day} de {MESES[hoy.month]} de {hoy.year}"
    )

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    if os.path.isfile(_LOGO_PATH):
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_after = Pt(6)
        logo_p.add_run().add_picture(_LOGO_PATH, width=Inches(2.2))

    doc.add_paragraph()

    p_fecha = doc.add_paragraph()
    p_fecha.paragraph_format.space_after = Pt(2)
    r_fecha = p_fecha.add_run(fecha_str)
    r_fecha.font.size = Pt(11)

    doc.add_paragraph()

    p_dest = doc.add_paragraph()
    p_dest.paragraph_format.space_after = Pt(0)
    r_dest = p_dest.add_run(
        "SEÑORES ADMINISTRACIÓN GENERAL DE PUERTOS"
    )
    r_dest.bold = True
    r_dest.font.size = Pt(11)

    p_pres = doc.add_paragraph()
    p_pres.paragraph_format.space_after = Pt(6)
    r_pres = p_pres.add_run("PRESENTE")
    r_pres.bold = True
    r_pres.font.size = Pt(11)

    doc.add_paragraph()

    mercaderia = op.mercaderia_a_mover or "mercadería"
    p_cuerpo = doc.add_paragraph()
    p_cuerpo.paragraph_format.space_after = Pt(4)
    r_cuerpo = p_cuerpo.add_run(
        f"Por medio de la presente se informa que con motivo "
        f"de la operación del buque {op.nombre_barco}, "
        f"se procederá al movimiento de {mercaderia}."
    )
    r_cuerpo.font.size = Pt(11)

    p_detalle = doc.add_paragraph()
    p_detalle.paragraph_format.space_after = Pt(8)
    r_detalle = p_detalle.add_run(
        "A continuación se detalla la nómina de "
        "transportes afectados a la operación:"
    )
    r_detalle.font.size = Pt(11)

    headers = [
        "Empresa",
        "Apellido y nombre",
        "DNI",
        "Camión",
        "Patente camión",
        "Patente acoplado",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    for a in op.asignaciones:
        row = table.add_row()
        vals = [
            a.empresa_snap,
            a.nombre_chofer_snap,
            a.dni_snap or "—",
            a.marca_camion_snap or "—",
            a.patente_camion_snap or "—",
            a.patente_acoplado_snap or "—",
        ]
        for i, v in enumerate(vals):
            cell = row.cells[i]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(v)
            run.font.size = Pt(10)

    doc.add_paragraph()

    firma_path = (
        Path(__file__).resolve().parent.parent.parent
        / "static" / "firmas" / "fernando_martinez.png"
    )
    if firma_path.is_file():
        firma_p = doc.add_paragraph()
        firma_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        firma_p.add_run().add_picture(
            str(firma_path), width=Inches(1.8)
        )

    p_nombre = doc.add_paragraph()
    p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre.paragraph_format.space_after = Pt(0)
    r_nombre = p_nombre.add_run("Fernando Martínez")
    r_nombre.bold = True
    r_nombre.font.size = Pt(11)

    p_cargo = doc.add_paragraph()
    p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cargo = p_cargo.add_run("MTR Logística")
    r_cargo.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = op.nombre_barco.replace(" ", "_").replace("/", "-")
    filename = f"carta_puerto_{safe_name}_{op_id}.docx"

    WORD_MIME = (
        "application/vnd.openxmlformats-"
        "officedocument.wordprocessingml.document"
    )
    return StreamingResponse(
        buf,
        media_type=WORD_MIME,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        },
    )


@router.post("/historial/{op_id}/asignar", response_class=HTMLResponse)
async def historial_asignar(
    op_id: int,
    nomina_id: int = Form(...),
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    op = db.query(mt.TransporteOperativo).filter(
        mt.TransporteOperativo.id == op_id,
        mt.TransporteOperativo.deleted_at == None,
    ).first()
    if not op:
        raise HTTPException(status_code=404, detail="Operativo no encontrado.")

    nomina = db.query(mt.TransporteNomina).filter(
        mt.TransporteNomina.id == nomina_id,
        mt.TransporteNomina.deleted_at == None,
    ).first()
    if not nomina:
        raise HTTPException(status_code=404, detail="Registro de nómina no encontrado.")

    # Evitar duplicados por nomina_id en el mismo operativo
    ya_asignado = db.query(mt.TransporteOperativoAsignacion).filter(
        mt.TransporteOperativoAsignacion.operativo_id == op_id,
        mt.TransporteOperativoAsignacion.nomina_id == nomina_id,
    ).first()
    if ya_asignado:
        return RedirectResponse(f"/transporte/historial/{op_id}", status_code=303)

    # Guardar snapshot de los datos actuales del transporte
    asignacion = mt.TransporteOperativoAsignacion(
        operativo_id          = op_id,
        nomina_id             = nomina_id,
        empresa_snap          = nomina.empresa,
        nombre_chofer_snap    = nomina.nombre_chofer,
        dni_snap              = nomina.dni,
        marca_camion_snap     = nomina.marca_camion,
        patente_camion_snap   = nomina.patente_camion,
        patente_acoplado_snap = nomina.patente_acoplado,
        assigned_by_id        = current_user.id,
    )
    db.add(asignacion)
    db.commit()
    return RedirectResponse(f"/transporte/historial/{op_id}", status_code=303)


@router.post("/historial/{op_id}/desasignar/{asig_id}")
async def historial_desasignar(
    op_id: int,
    asig_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    asig = db.query(mt.TransporteOperativoAsignacion).filter(
        mt.TransporteOperativoAsignacion.id == asig_id,
        mt.TransporteOperativoAsignacion.operativo_id == op_id,
    ).first()
    if not asig:
        raise HTTPException(status_code=404, detail="Asignación no encontrada.")

    db.delete(asig)
    db.commit()
    return RedirectResponse(f"/transporte/historial/{op_id}", status_code=303)


# ── EXPORTAR WORD ─────────────────────────────────────────────────────────────

# Ruta esperada del logo: static/logo_mtr.png en la raíz del proyecto.
_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
_LOGO_PATH    = os.path.join(_PROJECT_ROOT, "static", "logo_mtr.png")


@router.get("/historial/{op_id}/exportar-word")
async def historial_exportar_word(
    op_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    try:
        from docx import Document
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise HTTPException(
            status_code=501,
            detail="python-docx no está instalado. Ejecutá: pip install python-docx",
        )

    op = (
        db.query(mt.TransporteOperativo)
        .options(joinedload(mt.TransporteOperativo.asignaciones))
        .filter(
            mt.TransporteOperativo.id == op_id,
            mt.TransporteOperativo.deleted_at == None,
        )
        .first()
    )
    if not op:
        raise HTTPException(status_code=404, detail="Operativo no encontrado.")

    def _fmt_date(d):
        if not d:
            return "—"
        return d.strftime("%d/%m/%Y") if hasattr(d, "strftime") else str(d)

    def _remove_table_borders(table):
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "auto")
            tblBorders.append(b)
        tblPr.append(tblBorders)

    # ── Documento ──────────────────────────────────────────────────────────────
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Logo ───────────────────────────────────────────────────────────────────
    if os.path.isfile(_LOGO_PATH):
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_p.paragraph_format.space_after = Pt(6)
        logo_p.add_run().add_picture(_LOGO_PATH, width=Inches(2.2))

    # ── Título ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after  = Pt(14)
    r = title_p.add_run("Nómina de choferes y equipos")
    r.bold = True
    r.font.size = Pt(16)

    # ── Datos del operativo ────────────────────────────────────────────────────
    def _field(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        r_lbl = p.add_run(f"{label}: ")
        r_lbl.bold = True
        r_lbl.font.size = Pt(11)
        r_val = p.add_run(value or "—")
        r_val.font.size = Pt(11)

    _field("Buque",        op.nombre_barco)
    _field("Producto",     op.producto)
    _field("Cliente",      op.cliente)
    _field("Depósito",     op.deposito)
    _field("Fecha inicio", _fmt_date(op.fecha_inicio))
    _field("Fecha fin",    _fmt_date(op.fecha_fin))

    doc.add_paragraph()

    # ── Tabla de transportes ───────────────────────────────────────────────────
    headers = [
        "Empresa",
        "Apellido y nombre / Chofer",
        "DNI",
        "Camión",
        "Patente camión",
        "Patente acoplado",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    for a in op.asignaciones:
        row = table.add_row()
        vals = [
            a.empresa_snap,
            a.nombre_chofer_snap,
            a.dni_snap              or "—",
            a.marca_camion_snap     or "—",
            a.patente_camion_snap   or "—",
            a.patente_acoplado_snap or "—",
        ]
        for i, v in enumerate(vals):
            cell = row.cells[i]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(v)
            run.font.size = Pt(10)

    doc.add_paragraph()

    from pathlib import Path
    firma_path = (
        Path(__file__).resolve().parent.parent.parent
        / "static" / "firmas" / "fernando_martinez.png"
    )
    if firma_path.is_file():
        firma_p = doc.add_paragraph()
        firma_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        firma_p.add_run().add_picture(str(firma_path), width=Inches(1.8))

    p_nombre = doc.add_paragraph()
    p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre.paragraph_format.space_after = Pt(0)
    r_nombre = p_nombre.add_run("Fernando Martínez")
    r_nombre.bold = True
    r_nombre.font.size = Pt(11)

    p_cargo = doc.add_paragraph()
    p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cargo = p_cargo.add_run("MTR Logística")
    r_cargo.font.size = Pt(11)

    # ── Serializar ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = op.nombre_barco.replace(" ", "_").replace("/", "-")
    filename  = f"operativo_transporte_{safe_name}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

# ── NÓMINA PARA PUERTO ───────────────────────────────────────────────────────

@router.get("/historial/{op_id}/puerto", response_class=HTMLResponse)
async def historial_puerto(
    op_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    op = db.query(mt.TransporteOperativo).filter(
        mt.TransporteOperativo.id == op_id,
        mt.TransporteOperativo.deleted_at == None,
    ).first()
    if not op:
        raise HTTPException(status_code=404, detail="Operativo no encontrado.")

    excluidos_ids = {
        x.nomina_id
        for x in db.query(mt.TransportePuertoExclusion)
        .filter(mt.TransportePuertoExclusion.operativo_id == op_id)
        .all()
    }

    nomina = (
        db.query(mt.TransporteNomina)
        .filter(mt.TransporteNomina.deleted_at == None)
        .order_by(mt.TransporteNomina.empresa.asc(), mt.TransporteNomina.nombre_chofer.asc())
        .all()
    )

    incluidos = [n for n in nomina if n.id not in excluidos_ids]
    excluidos = [n for n in nomina if n.id in excluidos_ids]

    return templates.TemplateResponse(request, "transporte/historial_puerto.html", {
        "request": request,
        "user": current_user,
        "op": op,
        "incluidos": incluidos,
        "excluidos": excluidos,
    })


@router.post("/historial/{op_id}/puerto/quitar/{nomina_id}")
async def historial_puerto_quitar(
    op_id: int,
    nomina_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    existe = db.query(mt.TransportePuertoExclusion).filter(
        mt.TransportePuertoExclusion.operativo_id == op_id,
        mt.TransportePuertoExclusion.nomina_id == nomina_id,
    ).first()

    if not existe:
        db.add(mt.TransportePuertoExclusion(
            operativo_id=op_id,
            nomina_id=nomina_id,
            removed_by_id=current_user.id,
        ))
        db.commit()

    return RedirectResponse(f"/transporte/historial/{op_id}/puerto", status_code=303)


@router.post("/historial/{op_id}/puerto/restaurar/{nomina_id}")
async def historial_puerto_restaurar(
    op_id: int,
    nomina_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(_require_access),
):
    excl = db.query(mt.TransportePuertoExclusion).filter(
        mt.TransportePuertoExclusion.operativo_id == op_id,
        mt.TransportePuertoExclusion.nomina_id == nomina_id,
    ).first()

    if excl:
        db.delete(excl)
        db.commit()

    return RedirectResponse(f"/transporte/historial/{op_id}/puerto", status_code=303)
